

# %%
#########################################
# 1 워드문서 생성하기
#########################################


# 간단한 파일 생성
from docx import Document

# 새 문서 생성
doc = Document()

# 문단 추가
doc.add_paragraph("안녕하세요, Py!")

# 문서 저장
doc.save("예시.docx")
print("예시 파일이 생성되었습니다!")



# %%
#########################################
# 2 문단별로 폰트 종류와 폰트 크기 설정하기 
#########################################
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# 새 문서 생성
doc = Document()

# 첫 번째 문단 - 기본 텍스트
para1 = doc.add_paragraph("이 문장은 기본 스타일입니다.")

# 두 번째 문단 - 폰트와 크기 설정
para2 = doc.add_paragraph("이 문장은 폰트와 크기가 다릅니다.")
run = para2.runs[0]
run.font.name = "맑은고딕"
run.font.size = Pt(16)
run.bold = True
# run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')  # 한글 폰트 설정

# 저장
doc.save("예시.docx")
print("예시 파일이 생성되었습니다!")
# %%
#########################################
# 3 문단 정렬과 들여쓰기 
#########################################


from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm
# 새 문서 생성
doc = Document()
# 가운데 정렬된 문단 추가
para3 = doc.add_paragraph("이 문장은 가운데 정렬되었습니다.")
para3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

doc.add_paragraph("두번째 줄입니다....................................................")
# 들여쓰기가 설정된 문단 추가
para4 = doc.add_paragraph("이 문장은 들여쓰기가 설정되었습니다.")
para4_format = para4.paragraph_format
para4_format.left_indent = Cm(0.5)  # 왼쪽 들여쓰기

doc.save("예시.docx")
print("예시 파일이 생성되었습니다!")



# %%
#########################################
# 4 둘째 줄 들여쓰기 
#########################################
from docx import Document
from docx.shared import Cm

# 새 문서 생성
doc = Document()

# 매달린 들여쓰기가 적용된 문단 추가
para = doc.add_paragraph("이 문장은 둘째 줄 들여쓰기가 적용된 예제입니다. "
                         "첫 번째 줄은 들여쓰기가 없고, 나머지 줄은 들여쓰기되어 있습니다.")
para_format = para.paragraph_format
para_format.left_indent = Cm(1)  # 전체 들여쓰기
para_format.first_line_indent = Cm(-1)  # 첫 줄 내어쓰기

# 문서 저장
doc.save("예시.docx")
print("예시 파일이 생성되었습니다!")



# %%
#########################################
# 5 문단 앞/뒤 간격 설정 
#########################################
from docx import Document
from docx.shared import Pt

# 새 문서 생성
doc = Document()

doc.add_paragraph("첫번째문단입니다.......................................................................................")
# 위아래 여백이 설정된 문단 추가
para = doc.add_paragraph("이 문장은 위아래로 여백이 설정되었습니다.")
para_format = para.paragraph_format
para_format.space_before = Pt(30)  # 문단 위 여백 12 pt
para_format.space_after = Pt(18)   # 문단 아래 여백 18 pt
doc.add_paragraph("세번째문단입니다.......................................................................................")
# 문서 저장
doc.save("예시.docx")
print("예시.docx 파일이 생성되었습니다!")



# %%
#########################################
# 6 줄간격 설정 
#########################################
# 줄 간격이 설정된 문단 추가


# 새 문서 생성
doc = Document()
para2 = doc.add_paragraph("------------------------------------------------------------------------이 문장은 줄 간격이 설정되었습니다------------------------------------------------------------")
para2_format = para2.paragraph_format
para2_format.line_spacing = 1.1  # 줄 간격 1.5

# 문서 저장
doc.save("예시.docx")
print("예시 파일이 생성되었습니다!")

# %%
